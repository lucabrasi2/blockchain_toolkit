from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_BREAK
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


PROJECT_ROOT = Path.cwd()
OUTPUT = PROJECT_ROOT / "UBP_Complete_Source_Code.docx"

EXCLUDED_DIRS = {
    ".git",
    "venv",
    "node_modules",
    "__pycache__",
    ".pytest_cache",
    ".mypy_cache",
    ".ruff_cache",
    "dist",
    "build",
    "logs",
}

EXCLUDED_FILES = {
    ".env",
    ".env.local",
    ".env.production",
    "ubp.db",
    "blockchain_toolkit.zip",
}

SOURCE_EXTENSIONS = {
    ".py",
    ".js",
    ".jsx",
    ".ts",
    ".tsx",
    ".html",
    ".htm",
    ".css",
    ".scss",
    ".json",
    ".md",
    ".txt",
    ".sql",
    ".sh",
    ".yml",
    ".yaml",
    ".toml",
    ".ini",
    ".cfg",
    ".xml",
}


def is_source_file(path: Path) -> bool:
    if path.name in EXCLUDED_FILES:
        return False

    if path.suffix.lower() in SOURCE_EXTENSIONS:
        return True

    # Include extensionless executable/configuration files
    if path.name in {
        "Dockerfile",
        "Makefile",
        "Procfile",
        ".gitignore",
    }:
        return True

    return False


def configure_code_style(style):
    style.font.name = "Courier New"
    style._element.rPr.rFonts.set(
        qn("w:eastAsia"),
        "Courier New"
    )
    style.font.size = Pt(8)


def add_page_number(paragraph):
    paragraph.alignment = 2

    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"

    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


# ---------------------------------------------------------------------------
# Collect files
# ---------------------------------------------------------------------------

files = []

for path in PROJECT_ROOT.rglob("*"):
    if not path.is_file():
        continue

    relative = path.relative_to(PROJECT_ROOT)

    if any(part in EXCLUDED_DIRS for part in relative.parts):
        continue

    if is_source_file(path):
        files.append(path)

files.sort(key=lambda p: str(p.relative_to(PROJECT_ROOT)).lower())


# ---------------------------------------------------------------------------
# Create document
# ---------------------------------------------------------------------------

doc = Document()

# Title
title = doc.add_heading(
    "Universal Blockchain Platform (UBP)",
    level=0
)

subtitle = doc.add_paragraph()
subtitle.alignment = 1
run = subtitle.add_run("Complete Source Code Export")
run.bold = True
run.font.size = Pt(16)

info = doc.add_paragraph()
info.alignment = 1
info.add_run(
    f"Project: {PROJECT_ROOT}\n"
    f"Source files exported: {len(files)}\n"
)

doc.add_page_break()

# Table of contents-like index
doc.add_heading("Source Code Index", level=1)

for index, path in enumerate(files, start=1):
    relative = path.relative_to(PROJECT_ROOT)
    doc.add_paragraph(
        f"{index}. {relative}",
        style="List Number"
    )

doc.add_page_break()

# Code style
code_style = doc.styles["Normal"]
configure_code_style(code_style)


# ---------------------------------------------------------------------------
# Add each source file
# ---------------------------------------------------------------------------

for index, path in enumerate(files, start=1):

    relative = path.relative_to(PROJECT_ROOT)

    doc.add_heading(
        f"{index}. {relative}",
        level=1
    )

    metadata = doc.add_paragraph()
    metadata.add_run("Source path: ").bold = True
    metadata.add_run(str(relative))

    try:
        content = path.read_text(
            encoding="utf-8",
            errors="replace"
        )
    except Exception as exc:
        content = f"[Unable to read file: {exc}]"

    # Put source in a dedicated code paragraph.
    paragraph = doc.add_paragraph()
    paragraph.style = code_style

    for line in content.splitlines():
        paragraph.add_run(line)
        paragraph.add_run("\n")

    if index != len(files):
        doc.add_page_break()


# ---------------------------------------------------------------------------
# Footer
# ---------------------------------------------------------------------------

for section in doc.sections:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.text = "Universal Blockchain Platform (UBP) — Source Code Export"
    add_page_number(paragraph)


doc.save(OUTPUT)

print()
print("=" * 70)
print("UBP SOURCE EXPORT COMPLETE")
print("=" * 70)
print(f"Project:       {PROJECT_ROOT}")
print(f"Source files:  {len(files)}")
print(f"Word document: {OUTPUT}")
print(f"Size:           {OUTPUT.stat().st_size / (1024 * 1024):.2f} MB")
print("=" * 70)
